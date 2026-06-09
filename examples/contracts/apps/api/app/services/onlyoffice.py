"""OnlyOffice Document Server: конфиг DocEditor, нормализация jinja-тегов после
round-trip и помощники для callback.

Document Server редактирует master_skeleton.docx прямо в браузере. Главный риск —
теги docxtpl ({{ var }} / {% .. %}) при пересохранении в OnlyOffice могут оказаться
разорванными на несколько run'ов внутри абзаца, и тогда docxtpl их не распарсит.
normalize_docx_jinja_runs() склеивает такие run'ы обратно.
"""

from __future__ import annotations

import re
from io import BytesIO
from urllib.parse import urlsplit, urlunsplit

from docx import Document

from app.config import get_settings
from app.security import create_doc_download_token, onlyoffice_sign

settings = get_settings()

MASTER_CODE = "master_skeleton"

# Признак того, что в абзаце есть фрагмент jinja-тега (быстрый предфильтр).
_TAG_HINT = re.compile(r"{[{%#]|[}%#]}")
# Полный тег целиком (для проверки «не разорван ли»).
_FULL_TAG = re.compile(r"{{.*?}}|{%.*?%}|{#.*?#}", re.S)


def document_key(version: str) -> str:
    """Ключ документа для DS: уникален на версию контента (sha master_skeleton.docx).
    Требования DS: [A-Za-z0-9_-], длина ≤128. При смене контента ключ меняется →
    DS не отдаёт устаревшую закэшированную версию."""
    safe = re.sub(r"[^A-Za-z0-9_-]", "", version)[:120]
    return f"ms-{safe}"


def build_editor_config(*, version: str, user_id: int, user_name: str) -> dict:
    """Подписанный JWT конфиг DocEditor. document.url и callbackUrl — ВНУТРЕННИЕ
    (DS ходит к нашему api по docker-сети, не через публичный Traefik)."""
    base = settings.onlyoffice_internal_api_url.rstrip("/")
    dl_token = create_doc_download_token(MASTER_CODE)
    config: dict = {
        "documentType": "word",
        "document": {
            "title": "master_skeleton.docx",
            "fileType": "docx",
            "key": document_key(version),
            "url": f"{base}/api/templates/master-skeleton/raw?token={dl_token}",
            "permissions": {"edit": True, "download": True, "print": True, "review": False},
        },
        "editorConfig": {
            "mode": "edit",
            "lang": "ru",
            "callbackUrl": f"{base}/api/templates/master-skeleton/onlyoffice-callback",
            "user": {"id": str(user_id), "name": user_name},
            "customization": {
                "autosave": True,
                "forcesave": True,
                "compactHeader": False,
                "help": False,
            },
        },
    }
    config["token"] = onlyoffice_sign(config)
    return config


def rewrite_ds_download_url(url: str) -> str:
    """URL, который DS прислал в callback (payload.url), указывает на хост DS. Меняем
    scheme+host на внутренний адрес DS, чтобы api скачивал из кэша по docker-сети,
    а не через публичный Traefik/TLS (hairpin-NAT/loopback ненадёжен на одном VPS)."""
    internal = settings.onlyoffice_internal_ds_url.rstrip("/")
    if not internal:
        return url
    i = urlsplit(internal)
    u = urlsplit(url)
    return urlunsplit((i.scheme or u.scheme, i.netloc or u.netloc, u.path, u.query, u.fragment))


def _paragraph_needs_merge(runs: list) -> bool:
    """True, если в абзаце есть jinja-тег, который разорван между run'ами."""
    full = "".join(r.text for r in runs)
    if not _TAG_HINT.search(full):
        return False
    tags = _FULL_TAG.findall(full)
    if not tags:
        # есть «{{»/«{%», но целого тега нет → вероятно тег разорван (или незакрыт) → чиним
        return True
    # Если хотя бы один полный тег не лежит целиком ни в одном run — он разорван.
    return any(not any(tag in r.text for r in runs) for tag in tags)


def _merge_paragraph_runs(paragraph) -> None:
    """Склеивает текст всех run'ов абзаца в первый run (формат первого run сохраняется).
    Вызывается только для абзацев с разорванным тегом — потеря разбиения по run'ам
    безопасна, т.к. теги намеренно без внутреннего форматирования."""
    runs = paragraph.runs
    if len(runs) < 2:
        return
    runs[0].text = "".join(r.text for r in runs)
    for r in runs[1:]:
        r.text = ""


def _fix_paragraphs(paragraphs) -> None:
    for p in paragraphs:
        if _paragraph_needs_merge(p.runs):
            _merge_paragraph_runs(p)


def _fix_tables(tables) -> None:
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                _fix_paragraphs(cell.paragraphs)
                _fix_tables(cell.tables)


def normalize_docx_jinja_runs(data: bytes) -> bytes:
    """Проходит тело, таблицы и колонтитулы; чинит разорванные jinja-теги. Возвращает
    новый .docx в bytes. Если тегов нет — документ возвращается без изменений по смыслу."""
    doc = Document(BytesIO(data))

    _fix_paragraphs(doc.paragraphs)
    _fix_tables(doc.tables)
    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            _fix_paragraphs(hf.paragraphs)
            _fix_tables(hf.tables)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
