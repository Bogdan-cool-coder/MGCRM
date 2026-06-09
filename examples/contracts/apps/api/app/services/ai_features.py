"""Эпик 18 — AI Features: высокоуровневая бизнес-логика.

Этот модуль — слой между роутерами и `anthropic_client`. Здесь живут
prompt-builder'ы и parsers, специфичные для каждой фичи:
- `analyze_contract` — анализ контракта
- `prefill_deal` / `prefill_lead` — извлечение полей из истории активностей

Pure parts (build_*, parse_*, normalize_*) — pure-function, тестируются
без БД и без сети. Coordinated parts (analyze_*, prefill_*) — async,
работают с session, дёргают call_claude и log_ai_request.
"""
from __future__ import annotations

import logging
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select

from app.config import get_settings
from app.models import Activity, Contract
from app.services.anthropic_client import (
    CONTRACT_ANALYSIS_SYSTEM_PROMPT,
    DEAL_PREFILL_SYSTEM_PROMPT,
    AINotConfiguredError,
    AIResponseError,
    AIServiceError,
    ClaudeResponse,
    call_claude,
    log_ai_request,
    parse_json_response,
)

logger = logging.getLogger(__name__)


# ============ Whitelists ============

ALLOWED_ISSUE_SEVERITY: frozenset[str] = frozenset({"error", "warning", "info"})
ALLOWED_SECTION_STATUS: frozenset[str] = frozenset({"ok", "warning", "missing"})
ALLOWED_CONFIDENCE: frozenset[str] = frozenset({"high", "medium", "low"})

# Какие поля сделки/лида AI имеет право предлагать. Whitelist защищает фронт
# от того, чтобы Claude предложил `password` или `id`.
ALLOWED_DEAL_FIELDS: frozenset[str] = frozenset({
    "qualification_score",
    "expected_close_date",
    "amount",
    "currency",
    "description",
    "contact_name",
    "country",
    "company_name",
    "title",
    "lost_reason",
})

# RU-лейблы для UI (когда Claude не вернул label или вернул мусор).
DEFAULT_FIELD_LABELS: dict[str, str] = {
    "qualification_score": "Оценка квалификации",
    "expected_close_date": "Дата закрытия",
    "amount": "Сумма сделки",
    "currency": "Валюта",
    "description": "Описание потребности",
    "contact_name": "Контактное лицо",
    "country": "Страна",
    "company_name": "Компания",
    "title": "Название сделки",
    "lost_reason": "Причина проигрыша",
}


# ============ Contract analysis ============

# Кэш считается «свежим» если ai_analyzed_at < CACHE_TTL ago. Если force_refresh
# не передан и кэш свеж — возвращаем его без обращения к Claude.
CONTRACT_ANALYSIS_CACHE_TTL = timedelta(hours=1)


def build_contract_analysis_prompt(
    contract_number: str | None,
    contract_text: str,
    product_code: str,
    country_code: str,
) -> str:
    """Собирает user-промпт для анализа договора.

    contract_text — полный текст контракта (rendered, без шаблонных placeholder'ов).
    Truncate'аем до 50K символов чтобы влезть в context window и не пожирать
    бюджет. Если длиннее — добавляем подсказку про сокращение.
    """
    truncated = False
    if len(contract_text) > 50_000:
        contract_text = contract_text[:50_000]
        truncated = True

    parts = [
        f"Договор #{contract_number or '-'} ({product_code}, {country_code}):",
        "",
        contract_text,
    ]
    if truncated:
        parts.append("")
        parts.append("[Текст обрезан до 50000 символов — для анализа взяты первые секции]")
    return "\n".join(parts)


def parse_contract_analysis_response(raw: dict[str, Any]) -> dict[str, Any]:
    """Нормализует Claude-ответ в shape ContractAnalysisOut.

    Pure function: валидируем severity / status whitelist'ы, отбрасываем
    мусор, защищаемся от пропущенных ключей. Если Claude вернул не-список
    в issues — конвертируем в пустой список.
    """
    issues_raw = raw.get("issues") if isinstance(raw, dict) else None
    issues: list[dict[str, Any]] = []
    if isinstance(issues_raw, list):
        for item in issues_raw:
            if not isinstance(item, dict):
                continue
            severity = str(item.get("severity") or "info").lower()
            if severity not in ALLOWED_ISSUE_SEVERITY:
                severity = "info"
            issues.append({
                "severity": severity,
                "quote": str(item.get("quote") or "")[:2000],
                "explanation": str(item.get("explanation") or "").strip(),
                "suggestion": (
                    str(item["suggestion"]).strip()
                    if item.get("suggestion") is not None else None
                ),
                "section": (
                    str(item["section"]).strip()
                    if item.get("section") is not None else None
                ),
            })

    sections_raw = raw.get("standard_sections") if isinstance(raw, dict) else None
    sections: list[dict[str, str]] = []
    if isinstance(sections_raw, list):
        for item in sections_raw:
            if not isinstance(item, dict):
                continue
            status = str(item.get("status") or "ok").lower()
            if status not in ALLOWED_SECTION_STATUS:
                status = "ok"
            sections.append({
                "section": str(item.get("section") or "")[:200],
                "status": status,
            })

    recommendations_raw = raw.get("recommendations") if isinstance(raw, dict) else None
    recommendations: list[str] = []
    if isinstance(recommendations_raw, list):
        for item in recommendations_raw:
            if isinstance(item, str) and item.strip():
                recommendations.append(item.strip()[:1000])

    return {
        "issues": issues,
        "standard_sections": sections,
        "recommendations": recommendations,
    }


def _load_contract_text(contract: Contract) -> str:
    """Достаём текст контракта для анализа.

    Стратегия:
    1. Если есть docx_path — читаем .docx через python-docx (имеется в deps).
    2. Иначе — собираем из context (JSON) + master_skeleton.md (fallback).

    В MVP делаем простую конкатенацию параграфов; таблицы тоже включаем
    (для платежей/актов). Картинки/embedded — игнорируем.
    """
    if contract.docx_path and Path(contract.docx_path).exists():
        try:
            from docx import Document
            doc = Document(contract.docx_path)
            chunks: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    chunks.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        chunks.append(" | ".join(cells))
            return "\n".join(chunks)
        except Exception as e:  # noqa: BLE001
            logger.warning(
                "Failed to read docx %s: %s — falling back to context",
                contract.docx_path, e,
            )

    # Fallback: dump context-json в текст.
    import json
    return json.dumps(contract.context or {}, ensure_ascii=False, indent=2)


def is_analysis_cache_fresh(
    analyzed_at: datetime | None,
    *,
    ttl: timedelta = CONTRACT_ANALYSIS_CACHE_TTL,
    now: datetime | None = None,
) -> bool:
    """Pure-helper: свежий ли кэш AI-анализа."""
    if analyzed_at is None:
        return False
    current = now or datetime.now(UTC)
    # Если analyzed_at без tzinfo (старая запись) — считаем UTC.
    if analyzed_at.tzinfo is None:
        analyzed_at = analyzed_at.replace(tzinfo=UTC)
    return (current - analyzed_at) < ttl


async def analyze_contract(
    session: AsyncSession,
    contract: Contract,
    *,
    user_id: int,
    force_refresh: bool = False,
) -> dict[str, Any]:
    """Анализирует договор через Claude. Кэширует результат в contract.ai_analysis_json.

    Возвращает dict shape:
        {
          "issues": [...],
          "standard_sections": [...],
          "recommendations": [...],
          "model": str | None,
          "analyzed_at": datetime,
          "ai_tokens_used": int | None,
          "from_cache": bool,
        }

    Бросает:
    - AINotConfiguredError если ANTHROPIC_API_KEY пуст (caller → 503)
    - AIServiceError если Anthropic не отвечает (caller → 502)
    - AIResponseError если Claude вернул невалидный JSON (caller → 502)
    """
    # 1) Кэш: если свежий и force_refresh=False — возвращаем сразу.
    cached = contract.ai_analysis_json or {}
    if (
        not force_refresh
        and cached
        and is_analysis_cache_fresh(contract.ai_analyzed_at)
    ):
        return {
            "issues": cached.get("issues", []),
            "standard_sections": cached.get("standard_sections", []),
            "recommendations": cached.get("recommendations", []),
            "model": cached.get("model"),
            "analyzed_at": contract.ai_analyzed_at,
            "ai_tokens_used": cached.get("ai_tokens_used"),
            "from_cache": True,
        }

    # 2) Зовём Claude.
    contract_text = _load_contract_text(contract)
    prompt = build_contract_analysis_prompt(
        contract_number=contract.number,
        contract_text=contract_text,
        product_code=contract.product_code,
        country_code=contract.country_code,
    )

    response: ClaudeResponse | None = None
    status_str = "success"
    error_message: str | None = None
    try:
        response = await call_claude(
            prompt=prompt,
            system=CONTRACT_ANALYSIS_SYSTEM_PROMPT,
        )
        parsed_raw = parse_json_response(response.text)
        normalized = parse_contract_analysis_response(parsed_raw)
    except AINotConfiguredError:
        status_str = "not_configured"
        error_message = "ANTHROPIC_API_KEY not configured"
        await log_ai_request(
            session,
            user_id=user_id,
            feature="contract_analyze",
            entity_type="contract",
            entity_id=contract.id,
            status=status_str,
            error_message=error_message,
        )
        raise
    except AIResponseError as e:
        status_str = "error"
        error_message = str(e)[:500]
        await log_ai_request(
            session,
            user_id=user_id,
            feature="contract_analyze",
            entity_type="contract",
            entity_id=contract.id,
            status=status_str,
            response=response,
            error_message=error_message,
        )
        raise
    except AIServiceError as e:
        status_str = "error"
        error_message = str(e)[:500]
        await log_ai_request(
            session,
            user_id=user_id,
            feature="contract_analyze",
            entity_type="contract",
            entity_id=contract.id,
            status=status_str,
            error_message=error_message,
        )
        raise

    # 3) Кэшируем в contract.ai_analysis_json + ai_analyzed_at.
    now = datetime.now(UTC)
    contract.ai_analysis_json = {
        **normalized,
        "model": response.model,
        "ai_tokens_used": response.usage.total_tokens,
    }
    contract.ai_analyzed_at = now

    await log_ai_request(
        session,
        user_id=user_id,
        feature="contract_analyze",
        entity_type="contract",
        entity_id=contract.id,
        status=status_str,
        response=response,
    )

    return {
        **normalized,
        "model": response.model,
        "analyzed_at": now,
        "ai_tokens_used": response.usage.total_tokens,
        "from_cache": False,
    }


# ============ Deal/Lead Prefill ============

# Source whitelist для query-param. UI/CRUD передают `all` по дефолту.
ALLOWED_PREFILL_SOURCE: frozenset[str] = frozenset({
    "all", "tg", "email", "notes", "activities", "messages",
})


def normalize_source(source: str | None) -> str:
    """Нормализует source-параметр. Невалидный → 'all'."""
    if not source:
        return "all"
    s = source.lower().strip()
    if s in ALLOWED_PREFILL_SOURCE:
        return s
    return "all"


def normalize_period_days(period_days: int | None) -> int:
    """0 или None → 0 (all-time). Отрицательные → 0. >365 → 365."""
    if period_days is None or period_days <= 0:
        return 0
    return min(period_days, 365)


def build_prefill_prompt(
    entity_label: str,
    entity_id: int,
    activities_text: str,
) -> str:
    """Собирает user-промпт для prefill.

    entity_label: «сделке» / «лиду» (для естественного русского текста).
    activities_text: уже собранный текст активностей (через _format_activities).
    """
    if not activities_text.strip():
        return (
            f"Для {entity_label} #{entity_id} активностей не найдено. "
            "Верни suggestions: []."
        )
    return (
        f"Контекст из переписки / заметок / звонков по {entity_label} #{entity_id}:\n\n"
        f"{activities_text}\n\n"
        f"Извлеки BANT-данные и предложи значения полей. Верни JSON."
    )


def _format_activities(activities: list[Activity]) -> str:
    """Pure helper: форматирует список Activity в текст для промпта.

    Каждая активность — блок:
        [DATE] KIND (id=N): TITLE
        BODY (если есть)
    """
    if not activities:
        return ""
    chunks: list[str] = []
    for a in activities:
        date_str = (
            a.created_at.strftime("%Y-%m-%d") if a.created_at else "—"
        )
        title = (a.title or "").strip()
        line = f"[{date_str}] {a.kind} (id={a.id}): {title}"
        chunks.append(line)
        body = (a.body or "").strip()
        if body:
            # Truncate body чтобы не съесть весь контекст.
            chunks.append(body[:2000])
        chunks.append("")  # пустая строка между активностями
    return "\n".join(chunks)


async def _load_activities_for_target(
    session: AsyncSession,
    target_type: str,
    target_id: int,
    *,
    source: str,
    period_days: int,
) -> list[Activity]:
    """Достаём активности для prefill по фильтру source + period.

    source mapping:
    - 'all'        → все kind
    - 'tg'         → пока mapping в kind не реализован (Activity не различает
                      tg/email/wa). Берём kind='note' как прокси (заметки часто
                      содержат скопированную переписку). Этот mapping
                      эволюционирует с эпиком 5 (Inbox).
    - 'email'      → same — kind='note' (или kind='task' с note про email).
    - 'notes'      → kind='note' only.
    - 'activities' → kind in (call, meeting, task)
    - 'messages'   → kind='note'
    """
    stmt = select(Activity).where(
        Activity.target_type == target_type,
        Activity.target_id == target_id,
    )

    if source == "notes" or source == "messages" or source == "tg" or source == "email":
        stmt = stmt.where(Activity.kind == "note")
    elif source == "activities":
        stmt = stmt.where(Activity.kind.in_(("call", "meeting", "task")))
    # source == "all" → без фильтра

    if period_days > 0:
        threshold = datetime.now(UTC) - timedelta(days=period_days)
        stmt = stmt.where(Activity.created_at >= threshold)

    stmt = stmt.order_by(Activity.created_at.desc()).limit(50)

    rows = (await session.execute(stmt)).scalars().all()
    return list(rows)


def parse_prefill_response(raw: dict[str, Any]) -> dict[str, Any]:
    """Нормализует Claude-ответ в shape DealPrefillOut.suggestions.

    Pure function. Whitelist field'ов: только ALLOWED_DEAL_FIELDS попадают
    в результат. Whitelist confidence: high/medium/low (иначе → low).
    label дефолтится из DEFAULT_FIELD_LABELS если Claude не дал.
    """
    summary = ""
    if isinstance(raw, dict):
        summary_raw = raw.get("summary")
        if isinstance(summary_raw, str):
            summary = summary_raw.strip()[:500]

    suggestions_raw = raw.get("suggestions") if isinstance(raw, dict) else None
    suggestions: list[dict[str, Any]] = []
    if isinstance(suggestions_raw, list):
        for item in suggestions_raw:
            if not isinstance(item, dict):
                continue
            field = str(item.get("field") or "").strip()
            if field not in ALLOWED_DEAL_FIELDS:
                continue
            confidence = str(item.get("confidence") or "low").lower()
            if confidence not in ALLOWED_CONFIDENCE:
                confidence = "low"
            label = str(item.get("label") or "").strip()
            if not label:
                label = DEFAULT_FIELD_LABELS.get(field, field)
            suggestions.append({
                "field": field,
                "label": label,
                "suggested_value": item.get("suggested_value"),
                "confidence": confidence,
                "reasoning": str(item.get("reasoning") or "").strip()[:500],
                "source_activity_id": (
                    int(item["source_activity_id"])
                    if isinstance(item.get("source_activity_id"), int)
                    else None
                ),
                "source_text": (
                    str(item["source_text"]).strip()[:500]
                    if item.get("source_text") is not None
                    else None
                ),
            })
    return {"summary": summary, "suggestions": suggestions}


async def prefill_for_target(
    session: AsyncSession,
    target_type: str,
    target_id: int,
    *,
    user_id: int,
    source: str = "all",
    period_days: int = 30,
) -> dict[str, Any]:
    """Запускает AI-prefill по сделке/лиду.

    target_type: 'deal' | 'lead'. Возвращает dict:
        {
          "summary": str,
          "suggestions": list[dict],
          "model": str | None,
          "ai_tokens_used": int | None,
        }

    Бросает: те же исключения что analyze_contract.
    """
    if target_type not in ("deal", "lead"):
        raise ValueError(f"Unsupported target_type: {target_type!r}")

    normalized_source = normalize_source(source)
    normalized_period = normalize_period_days(period_days)

    activities = await _load_activities_for_target(
        session, target_type, target_id,
        source=normalized_source, period_days=normalized_period,
    )
    activities_text = _format_activities(activities)
    label = "сделке" if target_type == "deal" else "лиду"
    prompt = build_prefill_prompt(label, target_id, activities_text)

    feature_name = "deal_prefill" if target_type == "deal" else "lead_prefill"
    response: ClaudeResponse | None = None
    try:
        response = await call_claude(
            prompt=prompt,
            system=DEAL_PREFILL_SYSTEM_PROMPT,
        )
        parsed_raw = parse_json_response(response.text)
        normalized = parse_prefill_response(parsed_raw)
    except AINotConfiguredError:
        await log_ai_request(
            session, user_id=user_id, feature=feature_name,
            entity_type=target_type, entity_id=target_id,
            status="not_configured",
            error_message="ANTHROPIC_API_KEY not configured",
        )
        raise
    except (AIResponseError, AIServiceError) as e:
        await log_ai_request(
            session, user_id=user_id, feature=feature_name,
            entity_type=target_type, entity_id=target_id,
            status="error", response=response,
            error_message=str(e)[:500],
        )
        raise

    await log_ai_request(
        session, user_id=user_id, feature=feature_name,
        entity_type=target_type, entity_id=target_id,
        status="success", response=response,
    )
    return {
        **normalized,
        "model": response.model,
        "ai_tokens_used": response.usage.total_tokens,
    }


def get_ai_status() -> dict[str, Any]:
    """Pure helper для ping endpoint'а — если фронт захочет узнать
    включён ли AI. {available: bool, model: str | None}.
    """
    settings = get_settings()
    return {
        "available": bool(settings.anthropic_api_key),
        "model": settings.anthropic_model if settings.anthropic_api_key else None,
    }


__all__ = [
    "ALLOWED_ISSUE_SEVERITY",
    "ALLOWED_SECTION_STATUS",
    "ALLOWED_CONFIDENCE",
    "ALLOWED_DEAL_FIELDS",
    "ALLOWED_PREFILL_SOURCE",
    "CONTRACT_ANALYSIS_CACHE_TTL",
    "DEFAULT_FIELD_LABELS",
    "analyze_contract",
    "build_contract_analysis_prompt",
    "build_prefill_prompt",
    "get_ai_status",
    "is_analysis_cache_fresh",
    "normalize_period_days",
    "normalize_source",
    "parse_contract_analysis_response",
    "parse_prefill_response",
    "prefill_for_target",
]
