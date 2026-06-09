"""Программная сборка docxtpl-шаблонов финансовых документов (Ф5 / Фаза A).

Создаёт два .docx-шаблона с jinja2-разметкой docxtpl:
  - invoice.docx — счёт на оплату (инвойс): шапка (номер, дата, юрлицо+реквизиты,
    контрагент), таблица позиций (наименование/кол-во/цена/НДС/сумма), итоги
    (нетто/НДС/всего), сумма прописью, блок подписи+печати.
  - act.docx — акт выполненных работ: аналогичная структура, перечень работ.

Кладутся в templates/contracts_master/finance/ (этот каталог монтируется в контейнер
как /app/templates/contracts_master, см. docker-compose). doc_render.py резолвит их
оттуда (settings.templates_dir / "finance") с overlay-фолбэком в storage.

Контекст рендера (см. app/services/finance/doc_render.py::build_invoice_context):
  doc.number, doc.date, doc.due_date, doc.currency, doc.purpose,
  doc.amount_net, doc.vat_amount, doc.amount_gross, doc.amount_in_words,
  doc.signed (bool), doc.signer_name, doc.signed_date,
  seller.* (наше юрлицо: name, full_name, tax_id_label, tax_id, address, bank,
            bank_code_label, bank_code, account, director_short, director_position,
            phone, email),
  buyer.*  (контрагент: name, tax_id, address),
  lines[] (n, name, qty, unit_price, vat_rate, vat_amount, amount_gross),
  has_vat (bool — показывать колонку/строку НДС).

Запуск (из apps/api):
    .venv/bin/python scripts/build_finance_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Брендовый primary MACRO (#172747).
PRIMARY = RGBColor(0x17, 0x27, 0x47)


def _set_cell(cell, text: str, *, bold: bool = False, align=None, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = PRIMARY


def _build(kind: str) -> Document:
    """Собрать один шаблон. kind in {'invoice','act'}."""
    is_invoice = kind == "invoice"
    doc = Document()

    # Заголовок документа (jinja-условие в самом тексте недоступно для заголовка —
    # делаем статический заголовок по типу документа).
    if is_invoice:
        _heading(doc, "Счёт на оплату {{ doc.number }}")
        subtitle = "от {{ doc.date }}"
    else:
        _heading(doc, "Акт выполненных работ {{ doc.number }}")
        subtitle = "от {{ doc.date }}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle).font.size = Pt(11)

    doc.add_paragraph()

    # ── Стороны ──
    seller = doc.add_paragraph()
    seller.add_run("Исполнитель: ").bold = True
    seller.add_run(
        "{{ seller.full_name or seller.name }}, "
        "{{ seller.tax_id_label or 'ИНН' }} {{ seller.tax_id }}, "
        "{{ seller.address }}"
    )
    bank = doc.add_paragraph()
    bank.add_run(
        "{% if seller.bank or seller.account %}"
    )
    bank.add_run("Банк: ").bold = True
    bank.add_run(
        "{{ seller.bank }}, {{ seller.bank_code_label or 'БИК' }} {{ seller.bank_code }}, "
        "р/с {{ seller.account }}{% endif %}"
    )

    buyer = doc.add_paragraph()
    buyer.add_run("Заказчик: ").bold = True
    buyer.add_run(
        "{{ buyer.name }}{% if buyer.tax_id %}, ИНН/БИН {{ buyer.tax_id }}{% endif %}"
        "{% if buyer.address %}, {{ buyer.address }}{% endif %}"
    )

    if not is_invoice:
        period = doc.add_paragraph()
        period.add_run(
            "{% if doc.period %}Отчётный период: {{ doc.period }}{% endif %}"
        )

    doc.add_paragraph()

    # ── Таблица позиций ──
    headers = ["№", "Наименование", "Кол-во", "Цена", "НДС", "Сумма"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

    # Строка-цикл docxtpl. ВАЖНО (docxtpl 0.20.x): {%tr for%} и {%tr endfor%} должны
    # жить каждый в СВОЕЙ отдельной строке-маркере (вся строка удаляется директивой
    # {%tr%}), оборачивая ОДНУ data-строку. Поэтому 3 строки: [for][data][endfor].
    # Если for и endfor положить в одну строку — jinja падает 'unknown tag endfor'.
    for_row = table.add_row().cells
    _set_cell(for_row[0], "{%tr for line in lines %}", align=WD_ALIGN_PARAGRAPH.CENTER)

    row = table.add_row().cells
    _set_cell(row[0], "{{ line.n }}", align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell(row[1], "{{ line.name }}")
    _set_cell(row[2], "{{ line.qty }}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row[3], "{{ line.unit_price }}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row[4], "{{ line.vat_display }}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(row[5], "{{ line.amount_gross }}", align=WD_ALIGN_PARAGRAPH.RIGHT)

    end_row = table.add_row().cells
    _set_cell(end_row[0], "{%tr endfor %}", align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Итоги ──
    totals = doc.add_paragraph()
    totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals.add_run("Итого без НДС: {{ doc.amount_net }} {{ doc.currency }}\n")
    totals.add_run(
        "{% if has_vat %}НДС: {{ doc.vat_amount }} {{ doc.currency }}\n{% endif %}"
    )
    total_run = totals.add_run("Всего к оплате: {{ doc.amount_gross }} {{ doc.currency }}")
    total_run.bold = True

    words = doc.add_paragraph()
    words.add_run("Сумма прописью: ").bold = True
    words.add_run("{{ doc.amount_in_words }}")

    if is_invoice and True:
        purpose = doc.add_paragraph()
        purpose.add_run(
            "{% if doc.purpose %}Назначение: {{ doc.purpose }}{% endif %}"
        )

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Блок подписи / печати ──
    sig = doc.add_table(rows=2, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = sig.rows[0].cells[0]
    right = sig.rows[0].cells[1]
    role = "Исполнитель" if is_invoice else "Исполнитель"
    _set_cell(left, f"{role}:")
    _set_cell(right, "Заказчик:")

    left2 = sig.rows[1].cells[0]
    right2 = sig.rows[1].cells[1]
    # Подпись исполнителя: если документ подписан — печатаем ФИО подписанта и дату,
    # иначе оставляем линию для ручной подписи + место для печати (М.П.).
    _set_cell(
        left2,
        "{% if doc.signed %}{{ doc.signer_name }}\nДата: {{ doc.signed_date }}\n"
        "Подписано в системе MACRO CRM"
        "{% else %}_______________ / {{ seller.director_short }}\nМ.П.{% endif %}",
        size=9,
    )
    _set_cell(right2, "_______________ /\nМ.П.", size=9)

    return doc


def main() -> None:
    # parents[3] = repo root (scripts → apps/api → apps → repo root).
    out_dir = (
        Path(__file__).resolve().parents[3]
        / "templates"
        / "contracts_master"
        / "finance"
    )
    out_dir.mkdir(parents=True, exist_ok=True)

    for kind, fname in (("invoice", "invoice.docx"), ("act", "act.docx")):
        doc = _build(kind)
        path = out_dir / fname
        doc.save(str(path))
        print(f"Saved {kind} template → {path} ({path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
