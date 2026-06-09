"""Конвертирует оригинальный MacroCRM КЗ .docx в master_skeleton.docx
с jinja2-тегами для docxtpl.

Что делает:
1. Открывает _source_original.docx
2. По известному словарю замен подменяет жёстко закодированные значения
   на jinja2-теги (тексты идут построчно через все runs параграфов и ячеек)
3. Таблицы платежей и актирования превращает в шаблоны с {%tr for %}
4. Сохраняет как master_skeleton.docx

Запуск:
    python apps/api/scripts/build_skeleton_docx.py \\
        templates/contracts_master/_source_original.docx \\
        templates/contracts_master/master_skeleton.docx
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# ============= Словарь замен =============
# Жёстко закодированные строки оригинала → jinja2-теги
REPLACEMENTS: list[tuple[str, str]] = [
    # ----- Заголовок -----
    ("СУБЛИЦЕНЗИОННЫЙ  ДОГОВОР № ___", "СУБЛИЦЕНЗИОННЫЙ ДОГОВОР № {{ contract.number }}"),
    ("СУБЛИЦЕНЗИОННЫЙ ДОГОВОР № ___", "СУБЛИЦЕНЗИОННЫЙ ДОГОВОР № {{ contract.number }}"),

    # ----- Шапка: город и дата -----
    ("г. Астана", "г. {{ contract.city or country.default_city }}"),
    ("«ХХ» мая 2026 г.", "«{{ contract.date_day }}» {{ contract.date_month }} {{ contract.date_year }} г."),

    # ----- Стороны: реквизиты Лицензиара (Проптехсервис) -----
    ("Товарищество с ограниченной ответственностью «Проптехсервис Казахстан»",
     "{{ licensor.full_legal_form }} «{{ licensor.name }}»"),
    ("Амангельдиева Адильжана Канатұлы", "{{ licensor.director_genitive }}"),
    ("Амангельдиев А.К.", "{{ licensor.director_short }}"),
    ("Директора", "{{ licensor.director_position }}"),  # будет встречаться и для sublicensee — в шаблоне поправим
    # Реквизиты в таблице:
    ("ТОО «Проптехсервис Казахстан»", "{{ licensor.legal_form }} «{{ licensor.name }}»"),
    ("БИН: 221040026030", "{{ licensor.tax_id_label }}: {{ licensor.tax_id }}"),
    ("Казахстан, город Астана, район Сарыарка, улица Тарас Шевченко, здание 4/1, н.п. 17, почтовый индекс Z11A3X9",
     "{{ licensor.address }}"),
    ('АО "Банк ЦентрКредит"', "{{ licensor.bank }}"),
    ('АО «Банк ЦентрКредит»', "{{ licensor.bank }}"),
    ("БИК: KCJBKZKX", "{{ licensor.bank_code_label }}: {{ licensor.bank_code }}"),
    ("KZ 398562203126111204", "{{ licensor.account }}"),
    ("KZ398562203126111204", "{{ licensor.account }}"),
    ("info@proptechservice.kz", "{{ licensor.email }}"),
    ("proptechservice.kz", "{{ licensor.website }}"),

    # ----- Стороны: реквизиты Сублицензиата -----
    ("Общество с ограниченной ответственностью «___»",
     "{{ sublicensee.full_legal_form }} «{{ sublicensee.name }}»"),
    # В таблице — заменим отдельно через cell.text проход

    # ----- Подписи в Приложении 1 -----
    ("Ядыкин Б.К.", "{{ licensor.director_short }}"),
    ("Иванов И.И.", "{{ sublicensee.director_short }}"),

    # ----- Подписи в Приложении 2 -----
    ("/  Амангельдиев А.К.", "/ {{ licensor.director_short }}"),
    ("/ Амангельдиев А.К.", "/ {{ licensor.director_short }}"),

    # ----- Цена/срок (если есть в тексте оригинала) -----
    # (обычно в табл, обработаем там же)

    # ----- Продукт -----
    ("программа для ЭВМ «MacroCRM»", "программа для ЭВМ «{{ product.name }}»"),
    ("«MacroCRM»", "«{{ product.name }}»"),

    # ----- Сервер -----
    ("https://macroserver.kz/", "{{ country.server_url }}"),
    ("https://macroserver.kz", "{{ country.server_url }}"),

    # ----- Страна (ВАЖЕН порядок: родительный перед именительным) -----
    ("законодательством Республики Казахстан", "законодательством {{ country.name_genitive }}"),
    ("законодательства Республики Казахстан", "законодательства {{ country.name_genitive }}"),
    ("законодательству Республики Казахстан", "законодательству {{ country.name_genitive }}"),
    ("требований законов Республики Казахстан", "требований законов {{ country.name_genitive }}"),
    ("актов Республики Казахстан", "актов {{ country.name_genitive }}"),
    ("на территории Республики Казахстан", "на территории {{ country.name_genitive }}"),
    ("Республики Казахстан", "{{ country.name_genitive }}"),
    ("Республика Казахстан", "{{ country.name_full }}"),
]

# Замены для ячеек таблицы реквизитов Сублицензиата (правая колонка таблицы 1)
SUBLICENSEE_CELL_REPLACEMENTS: list[tuple[str, str]] = [
    # Шапка справа
    ("Сублицензиат:\nООО «_____________»",
     "Сублицензиат:\n{{ sublicensee.legal_form }} «{{ sublicensee.name }}»"),
    ("Сублицензиат:", "Сублицензиат:"),  # no-op, для якоря
    # БИН/реквизиты — отдельная замена потому что текст идентичен лицензиару
    # делаем уникальный prefix
]


def replace_in_paragraph(paragraph, replacements: list[tuple[str, str]]) -> bool:
    """Заменяет текст в параграфе на уровне всех <w:t> элементов (включая внутри
    гиперссылок, которые paragraph.runs не видит). Собираем full text из всех <w:t>,
    делаем замены, если изменилось — кладём результат в первый <w:t>, остальные чистим."""
    t_elements = list(paragraph._p.iter(qn("w:t")))
    if not t_elements:
        return False
    full = "".join(t.text or "" for t in t_elements)
    if not full:
        return False
    new = full
    for old, repl in replacements:
        if old in new:
            new = new.replace(old, repl)
    if new == full:
        return False
    # Кладём весь новый текст в первый <w:t>, остальные обнуляем
    t_elements[0].text = new
    # сохраняем пробелы
    t_elements[0].set(qn("xml:space"), "preserve")
    for t in t_elements[1:]:
        t.text = ""
    return True


def replace_in_cell(cell, replacements: list[tuple[str, str]]) -> int:
    """Применяет замены ко всем параграфам ячейки. Возвращает число изменённых параграфов."""
    n = 0
    for p in cell.paragraphs:
        if replace_in_paragraph(p, replacements):
            n += 1
    # Рекурсивно: ячейка может содержать вложенную таблицу
    for nested in cell.tables:
        for row in nested.rows:
            for c in row.cells:
                n += replace_in_cell(c, replacements)
    return n


def replace_globally(doc: Document, replacements: list[tuple[str, str]]) -> int:
    """Все параграфы документа + все ячейки всех таблиц + header/footer."""
    n = 0
    for p in doc.paragraphs:
        if replace_in_paragraph(p, replacements):
            n += 1
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                n += replace_in_cell(c, replacements)
    for section in doc.sections:
        for p in section.header.paragraphs:
            if replace_in_paragraph(p, replacements):
                n += 1
        for p in section.footer.paragraphs:
            if replace_in_paragraph(p, replacements):
                n += 1
    return n


def adapt_sublicensee_requisites(doc: Document) -> None:
    """В таблице реквизитов (Table 1) правая колонка содержит копию реквизитов лицензиара
    или заглушки. Меняем на jinja-теги sublicensee."""
    if len(doc.tables) < 2:
        return
    table = doc.tables[1]
    if len(table.columns) < 3:
        return
    # Колонка 2 (третья) — сублицензиат
    SUB_REPL = [
        ("Сублицензиат:\nООО «_____________»",
         "Сублицензиат:\n{{ sublicensee.legal_form }} «{{ sublicensee.name }}»"),
        ("Сублицензиат:\nООО «{{ sublicensee.legal_form }} «{{ sublicensee.name }}»»",  # на случай если первая replace уже частично применилась
         "Сублицензиат:\n{{ sublicensee.legal_form }} «{{ sublicensee.name }}»"),
        # Реквизиты сублицензиата — заменяем готовый licensor-блок (он мог быть скопирован)
        # на пустой шаблон с jinja-тегами
    ]
    # Для каждой ячейки правой колонки полностью переписываем содержимое
    SUBLICENSEE_BODY_TEMPLATE = (
        "{{ sublicensee.tax_id_label or country.tax_id_label }}: {{ sublicensee.tax_id }}\n"
        "Адрес: {{ sublicensee.address }}\n"
        "Банк: {{ sublicensee.bank }}\n"
        "{{ sublicensee.bank_code_label or country.bank_code_label }}: {{ sublicensee.bank_code }}\n"
        "Номер счёта: {{ sublicensee.account }}\n"
        "Эл. адрес: {{ sublicensee.email }}\n"
        "Вебсайт: {{ sublicensee.website }}"
    )

    for ri, row in enumerate(table.rows):
        if len(row.cells) < 3:
            continue
        right_cell = row.cells[2]
        text = right_cell.text.strip()
        if ri == 0:
            # Шапка: "Сублицензиат: ..."
            for p in right_cell.paragraphs:
                replace_in_paragraph(p, SUB_REPL)
        elif ri == 1:
            # Тело реквизитов: переписываем
            for p in right_cell.paragraphs:
                p.clear()
            first_p = right_cell.paragraphs[0] if right_cell.paragraphs else right_cell.add_paragraph()
            first_p.text = SUBLICENSEE_BODY_TEMPLATE
        elif ri == 2:
            # Подпись: "Директор\nИванов И.И. _______________" → переменные
            for p in right_cell.paragraphs:
                replace_in_paragraph(p, [
                    ("Иванов И.И.", "{{ sublicensee.director_short }}"),
                    ("Директор\n", "{{ sublicensee.director_position }}\n"),
                ])


def adapt_payment_table(doc: Document) -> None:
    """Таблица 3: график платежей. Шаблон с 3 строками: open / content / close.
    docxtpl при рендере удаляет строки с {%tr for/endfor %}, размножает content по items."""
    if len(doc.tables) < 4:
        return
    table = doc.tables[3]
    if len(table.rows) < 2 or len(table.columns) < 6:
        return

    # Удаляем все строки кроме header
    rows_to_delete = list(table.rows)[1:]
    for r in rows_to_delete:
        r._element.getparent().remove(r._element)

    # Строка 1: open
    open_row = table.add_row()
    open_row.cells[0].text = "{%tr for payment in license.payment_schedule %}"
    # Строка 2: content
    content_row = table.add_row()
    cells = content_row.cells
    cells[0].text = "{{ payment.number }}"
    cells[1].text = "{{ payment.amount }}"
    cells[2].text = "{{ payment.vat }}"
    cells[3].text = "{{ payment.due_date }}"
    cells[4].text = "{{ payment.period }}"
    cells[5].text = "{{ payment.included }}"
    # Строка 3: close
    close_row = table.add_row()
    close_row.cells[0].text = "{%tr endfor %}"


def adapt_act_table(doc: Document) -> None:
    """Таблица 4: график актирования. Аналогично платежам — 3 строки шаблона."""
    if len(doc.tables) < 5:
        return
    table = doc.tables[4]
    if len(table.rows) < 2 or len(table.columns) < 5:
        return

    rows_to_delete = list(table.rows)[1:]
    for r in rows_to_delete:
        r._element.getparent().remove(r._element)

    open_row = table.add_row()
    open_row.cells[0].text = "{%tr for act in license.act_schedule %}"
    content_row = table.add_row()
    cells = content_row.cells
    cells[0].text = "{{ act.number }}"
    cells[1].text = "{{ act.amount }}"
    cells[2].text = "{{ act.vat }}"
    cells[3].text = "{{ act.sign_by_date }}"
    cells[4].text = "{{ act.period }}"
    close_row = table.add_row()
    close_row.cells[0].text = "{%tr endfor %}"


def adapt_modules_table(doc: Document) -> None:
    """Таблица 2: спецификация модулей. Делаем динамической по product.modules_flat.
    Структура: header + строки. Заменяем строки данных на шаблонную с {%tr for %}."""
    if len(doc.tables) < 3:
        return
    table = doc.tables[2]
    if len(table.rows) < 2 or len(table.columns) < 3:
        return

    rows_to_delete = list(table.rows)[1:]
    for r in rows_to_delete:
        r._element.getparent().remove(r._element)

    open_row = table.add_row()
    open_row.cells[0].text = "{%tr for m in product.modules_flat %}"
    content_row = table.add_row()
    content_row.cells[0].text = "{{ m.section }}"
    content_row.cells[1].text = "{{ m.name }}"
    content_row.cells[2].text = "{{ m.status }}"
    close_row = table.add_row()
    close_row.cells[0].text = "{%tr endfor %}"


def adapt_brief_tables(doc: Document) -> None:
    """Таблицы 5 (бриф) и 6 (тех.параметры). Делаем динамическими по
    product.brief_fields / product.tech_fields. Колонка «Содержание» — пустая
    (менеджер дозаполняет в Word после генерации, либо позже через карточку)."""
    configs = [
        (5, "product.brief_fields"),
        (6, "product.tech_fields"),
    ]
    for table_idx, loop_var in configs:
        if len(doc.tables) <= table_idx:
            continue
        table = doc.tables[table_idx]
        if len(table.columns) < 3:
            continue
        rows_to_delete = list(table.rows)[1:]
        for r in rows_to_delete:
            r._element.getparent().remove(r._element)

        item = "f"
        open_row = table.add_row()
        open_row.cells[0].text = f"{{%tr for {item} in {loop_var} %}}"
        content_row = table.add_row()
        content_row.cells[0].text = f"{{{{ {item}.id }}}}"
        content_row.cells[1].text = f"{{{{ {item}.label }}}}"
        content_row.cells[2].text = ""  # содержание — заполняется вручную
        close_row = table.add_row()
        close_row.cells[0].text = "{%tr endfor %}"


def adapt_procedure_table(doc: Document) -> None:
    """Таблица 7: процедура настройки ПО. Только подставляем дату начала внедрения."""
    if len(doc.tables) < 8:
        return
    table = doc.tables[7]
    PROC_REPL = [
        ("Плановая дата начала настройки ПО: 24.01.2025 г.",
         "Плановая дата начала настройки ПО: {{ license.implementation_start_date }}"),
    ]
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, PROC_REPL)


def adapt_license_specification(doc: Document) -> None:
    """В Приложении 1 нужно подставить тип лицензии, срок, стоимость.
    Эти поля идут в обычном параграфе, не в таблице."""
    # Тип лицензии: "Тип лицензии: Стандартная"
    # Срок: "с 01.01.2026 до 01.01.2027 (12 мес)"
    # Стоимость: "100 000 (Сто тысяч) тенге"
    SPEC_REPL = [
        ("Тип лицензии: Стандартная", "Тип лицензии: {{ license.type }}"),
        # Срок предоставления — есть в тексте
    ]
    for p in doc.paragraphs:
        replace_in_paragraph(p, SPEC_REPL)


def main(src: Path, dst: Path) -> None:
    print(f"Открываю {src}…")
    doc = Document(str(src))

    print("Применяю глобальные замены…")
    n = replace_globally(doc, REPLACEMENTS)
    print(f"  заменено в {n} параграфах")

    print("Адаптирую реквизиты Сублицензиата (Table 1)…")
    adapt_sublicensee_requisites(doc)

    print("Адаптирую таблицу платежей (Table 3) — превращаю в {%tr for %}…")
    adapt_payment_table(doc)

    print("Адаптирую таблицу актирования (Table 4) — превращаю в {%tr for %}…")
    adapt_act_table(doc)

    print("Делаю таблицу модулей (Table 2) динамической по product.modules_flat…")
    adapt_modules_table(doc)

    print("Делаю бриф/тех.параметры (Tables 5, 6) динамическими…")
    adapt_brief_tables(doc)

    print("Подставляю дату внедрения в Table 7 (процедура настройки)…")
    adapt_procedure_table(doc)

    print("Подставляю тип лицензии в Приложение 1…")
    adapt_license_specification(doc)

    doc.save(str(dst))
    size_kb = dst.stat().st_size / 1024
    print(f"\n✅ Сохранено: {dst} ({size_kb:.1f} КБ)")
    print()
    print("Что нужно проверить вручную (открыв в Word):")
    print("  1. Заголовок 'СУБЛИЦЕНЗИОННЫЙ ДОГОВОР № {{ contract.number }}' — выровнен")
    print("  2. Таблица реквизитов (Table 1): слева ТОО Проптехсервис → licensor, справа sublicensee")
    print("  3. Таблица платежей (Table 3): должна остаться одна строка с {%tr for payment ... %}")
    print("  4. Таблица актирования (Table 4): то же")
    print("  5. Бриф и тех.параметры (Tables 5, 6): правая колонка пустая для дозаполнения")
    print("  6. В Приложении 1 параграф 'Тип лицензии: {{ license.type }}'")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python build_skeleton_docx.py <source.docx> <master_skeleton.docx>")
        sys.exit(1)
    main(Path(sys.argv[1]), Path(sys.argv[2]))
