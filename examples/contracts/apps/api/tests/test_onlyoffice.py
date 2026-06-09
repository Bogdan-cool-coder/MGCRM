"""OnlyOffice: нормализация разорванных jinja-тегов + конфиг DocEditor (без БД).

Главный риск Фазы 5 — OnlyOffice при пересохранении может разбить тег docxtpl
({{ x }} / {%tr .. %}) на несколько run'ов, и docxtpl его не распарсит.
normalize_docx_jinja_runs() должен это чинить, не ломая нетронутые теги и текст.
"""

from io import BytesIO

from docx import Document
from docxtpl import DocxTemplate

from app.config import get_settings
from app.security import onlyoffice_verify
from app.services.onlyoffice import (
    build_editor_config,
    document_key,
    normalize_docx_jinja_runs,
    rewrite_ds_download_url,
)


def _docx(make) -> bytes:
    doc = Document()
    make(doc)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_normalize_merges_split_inline_tag():
    def make(doc):
        p = doc.add_paragraph()
        p.add_run("Договор № ")
        p.add_run("{{ ")
        p.add_run("contract.number")
        p.add_run(" }}")

    p = Document(BytesIO(normalize_docx_jinja_runs(_docx(make)))).paragraphs[0]
    assert "{{ contract.number }}" in p.text
    assert any("{{ contract.number }}" in r.text for r in p.runs)


def test_normalized_split_tag_renders_via_docxtpl():
    def make(doc):
        p = doc.add_paragraph()
        p.add_run("{{")
        p.add_run(" contract.number ")
        p.add_run("}}")

    tpl = DocxTemplate(BytesIO(normalize_docx_jinja_runs(_docx(make))))
    tpl.render({"contract": {"number": "Д-42"}})
    out = BytesIO()
    tpl.save(out)
    assert "Д-42" in Document(BytesIO(out.getvalue())).paragraphs[0].text


def test_normalize_keeps_intact_tag_and_formatting():
    # тег целиком в одном run рядом с жирным текстом — не трогаем, формат сохраняется
    def make(doc):
        p = doc.add_paragraph()
        r1 = p.add_run("ВАЖНО")
        r1.bold = True
        p.add_run(" {{ contract.number }}")

    p = Document(BytesIO(normalize_docx_jinja_runs(_docx(make)))).paragraphs[0]
    assert "{{ contract.number }}" in p.text
    assert any(r.text == "ВАЖНО" and r.bold for r in p.runs)


def test_normalize_fixes_split_tr_tag_in_table():
    def make(doc):
        t = doc.add_table(rows=1, cols=1)
        p = t.rows[0].cells[0].paragraphs[0]
        p.add_run("{%tr ")
        p.add_run("for x in items")
        p.add_run(" %}")

    cell = Document(BytesIO(normalize_docx_jinja_runs(_docx(make)))).tables[0].rows[0].cells[0]
    assert "{%tr for x in items %}" in cell.paragraphs[0].text


def test_normalize_preserves_plaintext():
    def make(doc):
        p = doc.add_paragraph()
        p.add_run("Просто ")
        p.add_run("текст без тегов")

    assert Document(BytesIO(normalize_docx_jinja_runs(_docx(make)))).paragraphs[0].text == "Просто текст без тегов"


def test_document_key_is_safe_and_bounded():
    assert document_key("abc123def456") == "ms-abc123def456"
    assert document_key("a/b c:d") == "ms-abcd"
    long_key = document_key("x" * 500)
    assert long_key.startswith("ms-") and len(long_key) <= 124


def test_rewrite_ds_download_url_swaps_host():
    get_settings().onlyoffice_internal_ds_url = "http://onlyoffice"
    out = rewrite_ds_download_url("https://office.example.com/cache/files/abc/output.docx?x=1")
    assert out == "http://onlyoffice/cache/files/abc/output.docx?x=1"


def test_build_editor_config_signs_and_targets_internal_api():
    s = get_settings()
    s.onlyoffice_jwt_secret = "test-secret"
    s.onlyoffice_internal_api_url = "http://api:8000"
    config = build_editor_config(version="deadbeef", user_id=7, user_name="Юрист")
    assert config["documentType"] == "word"
    assert config["document"]["key"] == "ms-deadbeef"
    assert config["document"]["url"].startswith(
        "http://api:8000/api/templates/master-skeleton/raw?token="
    )
    assert (
        config["editorConfig"]["callbackUrl"]
        == "http://api:8000/api/templates/master-skeleton/onlyoffice-callback"
    )
    claims = onlyoffice_verify(config["token"])
    assert claims is not None and claims["document"]["key"] == "ms-deadbeef"
